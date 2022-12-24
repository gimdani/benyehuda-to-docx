from bs4 import BeautifulSoup
from urllib.request import Request, urlopen
import requests
import re
from selenium import webdriver

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

import tkinter.tix as tix

from string import punctuation

none = '-!-!-'
headlines_dict = {
    'שירה': 'works-poetry',
    'פרוזה': 'works-prose',
    'מחזות': 'works-drama',
    'מאמרים ומסות': 'works-article',
    'מכתבים': 'works-letters',
    'עיון': 'works-reference',
    'יצירות מתורגמות': 'works-translations',
    'זכרונות ויומנים': 'works-memoir'

}



def get_Urls(url_lxml, tag1_id="", tag2_txt="", tagType="", t=0, penetrate=0):
    
    soup = url_lxml

    tags = []
    
    if penetrate==0:
        soups = soup.find(tagType, {'id': tag1_id}).next_siblings
    else:
        soups = (soup.find(tagType, {'id': tag1_id}).findChildren('a'))
        
    for tag in soups: 
        if tag.text == tag2_txt:
            break
        elif not tag =='\n':
            tags.append(tag)
            #print (tag)
        if tag2_txt==none and not tagType=='div' and t==0:
            if tag.name=='h3' or tag.name=='h4':
                break
    res={}
    links=[]
    for link in tags:
        found =''
        
        if str(link).__contains__("h4"):
            found = re.search('>(.+?)<', str(link)).group(1)
        elif t==0 and str(link).__contains__("h5"):
            found = re.search('>(.+?)<', str(link)).group(1)
        if t == 0 and found=='':
            try:
                found = re.search('"(.+?)"', str(link)).group(1)
            except AttributeError:
                tyhbbgdhujnjy=1
        if found=='' and t==0:
            try:
                found = re.search('<strong>(.+?)</strong>', str(link)).group(1)
            except AttributeError:
                if str(link) =='<br>':
                    found="end_of_tab" 
                else:
                    link=''      
        if not link=='' and not found=='':
            if(t==0):
                links.append(found)
            else:
                res[found]= re.search('"(.+?)"', str(link)).group(1)
            #print(found)
    if t==0:
        return links
    return res


def get_Titles_L1(url_html):

    soup = url_html

    headlines = soup.find(class_= 'mainlist').find_all(class_='headline-1-v02')
    res=[]
    for x in headlines:
        #print(x.text.strip())
        res.append(x.text.strip())
    return res

def get_Titles_L2(url_html, header, level):
    headlines = {}
    
    soup = url_html

    #soup = BeautifulSoup(html)
    
    divTag = soup.find_all("div", {'id': header})

    for tag in divTag:
        tdTags = tag.find_all(level)
        for tag in tdTags:
            headlines[tag.text] = tag.get('id')

    return headlines

def extract_text(url):
    
    # Assign URL
    
    # Fetch raw HTML content
    html_content = requests.get(url).text
    
    # Now that the content is ready, iterate
    # through the content using BeautifulSoup:
    soup = BeautifulSoup(html_content, "html.parser")
    
    # similarly to get all the occurrences of a given tag
    res = soup.find("div", class_="work-title name-under-btn").text.split('\n')
    res = list(filter(lambda x: x != '', res))
    text = soup.find("div", {'id':'actualtext'}).text
    res.append(text.replace('\n\n', '\n').replace('\n\n','\n'))
    return res

def write_txt(doc, text, heading_level, author=False):
    heading = text[0]
    if author:
        heading += ' / ' + text[1] 
    h = doc.add_heading(heading, level=int(heading_level))
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = doc.add_paragraph(text[2])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    print("writing: "+ heading + ' / ' + text[1] )

def hebrew_text(text):
    text = text.split()[::-1]
    for i in text:
        if i[-1] in punctuation:
            text[text.index(i)] = i[-1]+i[0:-1]
    return " ".join(text)    

class View(object):

        works_new = {}

        def __init__(self, root):
            self.root = root
            self.root.geometry('370x550')
            self.makeCheckList()
        def makeCheckList(self):
            turn_off = tix.Button(self.root, text="submit", command=self.submit)
            turn_off.pack()
            self.root.cl = tix.CheckList(self.root, browsecmd=self.selectItem)
            self.root.cl.pack(expand=True,fill='both')
            self.root.cl.hlist.add('works', text = 'יצירות')
            self.root.cl.setstatus('works', 'off')
            for t1 in works:
                self.root.cl.hlist.add('works.'+t1.replace('.',''), text = hebrew_text(t1))
                self.root.cl.setstatus('works.'+t1.replace('.',''), 'off')
                if isinstance(works[t1], dict):
                    for t2 in works[t1]:
                        if not t2 == none:
                            self.root.cl.hlist.add('works.'+t1.replace('.','')+'.'+t2.replace('.',''), text = hebrew_text(t2))
                            self.root.cl.setstatus('works.'+t1.replace('.','')+'.'+t2.replace('.',''), 'off')
                            if isinstance(works[t1][t2], dict): 
                                for t3 in works[t1][t2]:                         
                                    if not t3 == none:
                                        self.root.cl.hlist.add('works.'+t1.replace('.','')+'.'+t2.replace('.','')+'.'+t3.replace('.',''), text = hebrew_text(t3))
                                        self.root.cl.setstatus('works.'+t1.replace('.','')+'.'+t2.replace('.','')+'.'+t3.replace('.',''), 'off')
                                        if isinstance(works[t1][t2][t3], dict): 
                                            for t4 in works[t1][t2][t3]:                         
                                                if not t4 == none:
                                                    self.root.cl.hlist.add('works.'+t1.replace('.','')+'.'+t2.replace('.','')+'.'+t3.replace('.','')+'.'+t4.replace('.',''), text = hebrew_text(t4))
                                                    self.root.cl.setstatus('works.'+t1.replace('.','')+'.'+t2.replace('.','')+'.'+t3.replace('.','')+'.'+t4.replace('.',''), 'off')
            #self.root.cl.autosetmode()
        def selectItem(self, item):
            if self.root.cl.getstatus(item) == 'on':
                self.autoCheckChildren(item, True)
            if self.root.cl.getstatus(item) == 'off':
                self.autoCheckChildren(item, False)
        
        def autoCheckChildren(self, i_item, stat):
            item = i_item
            if stat:
                if self.root.cl.hlist.info_children(item):
                    for child in self.root.cl.hlist.info_children(item):
                        self.root.cl.setstatus(child, "on")
                        self.autoCheckChildren(child, True)
            elif not stat:
                if self.root.cl.hlist.info_children(item):
                    for child in self.root.cl.hlist.info_children(item):
                        self.root.cl.setstatus(child, "off")
                        self.autoCheckChildren(child, False)
        
        def submit(self):
            self.works_new = self.make_works_new(works, 'works')
            self.root.quit()
        
        def make_works_new(self,works, name):
            tmp = {}
            for t in works:
                if isinstance(works[t], dict):
                    tmp[t]=self.make_works_new(works[t],name+'.'+t.replace('.',''))
                elif not t == none:
                    if self.root.cl.getstatus(name+'.'+t.replace('.','')) == 'on':
                        tmp[t]=works[t]

            return tmp



        
def choices_UI(works):
                
    root = tix.Tk()
    view = View(root)
    root.update()
    root.mainloop()
    return view.works_new
  
    
# add comments!!! 

if __name__ == '__main__':
    url = input("insert url (example: https://benyehuda.org/author/20): ")
    print ("Wait a sec")

    html_page = urlopen(url)
    url_lxml = BeautifulSoup(html_page, "lxml")

    response = requests.get(url)
    url_html = BeautifulSoup(response.text, 'html.parser')

    print ("Loading data. This might take a while...")
    works ={}
    titles1 = get_Titles_L1(url_html)
    for title1 in titles1:
        works[title1] = get_Titles_L2(url_html, headlines_dict[title1], 'h3' )

        if not works[title1]=={}:
            work = works[title1]
            for i in range(len(work)):
                work[none] = ''
                tmp = get_Urls(url_lxml, list(work.values())[i], list(work.keys())[i+1], 'h3', 1)
                if not tmp=={}:
                    tmp[none] = ''
                    work[list(work.keys())[i]] = tmp

    for t1 in works:
        if works[t1] =={}:
            works[t1]=get_Urls(url_lxml, headlines_dict[t1], none, 'div', penetrate=1)
        else:
            for t2 in works[t1]: #wrong!!!
                if not isinstance(works[t1][t2], dict) and not t2==none:
                    tmp = get_Urls(url_lxml, works[t1][t2], list(works[t1])[list(works[t1]).index(t2)+1], 'h3', penetrate=1)
                    tmp += (get_Urls(url_lxml, works[t1][t2], list(works[t1])[list(works[t1]).index(t2)+1], 'h3'))
                    works[t1][t2] = tmp
                elif isinstance(works[t1][t2], dict):
                    for t3 in works[t1][t2]:
                        if not t3==none:
                            works[t1][t2][t3]=get_Urls(url_lxml, works[t1][t2][t3], list(works[t1][t2])[list(works[t1][t2]).index(t3)+1], 'h4')
                
    

    works_new = choices_UI(works)

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'David'
    font.size = Pt(14)
    font.rtl = True

    for t1 in works_new:
        if not works_new[t1]=={} and not works_new[t1]==[]:
            h = document.add_heading(t1, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if isinstance(works_new[t1], list):
                for item_url in works_new[t1]:
                    if str(item_url).__contains__("read"):
                        write_txt(document, extract_text(item_url), 3)
                    else:
                        h = document.add_heading(item_url, level=2)
                        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif isinstance(works_new[t1], dict):
                for t2 in works_new[t1]:
                    if not works_new[t1][t2]=={} and not works_new[t1][t2]==[]:
                        h = document.add_heading(t2, level=2)
                        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        if isinstance(works_new[t1][t2], list):
                            for item_url in works_new[t1][t2]:
                                if str(item_url).__contains__("read"):
                                    write_txt(document, extract_text(item_url), 4)
                                else:
                                    h = document.add_heading(item_url, level=3)
                                    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif isinstance(works_new[t1][t2], dict):
                            for t3 in works_new[t1][t2]:
                                if not works_new[t1][t2][t3]=={} and not works_new[t1][t2][t3]==[]:
                                    h = document.add_heading(t3, level=3)
                                    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                    if isinstance(works_new[t1][t2][t3], list):
                                        for item_url in works_new[t1][t2][t3]:
                                            if str(item_url).__contains__("read"):
                                                write_txt(document, extract_text(item_url), 5)
                                            else:
                                                h = document.add_heading(item_url, level=4)
                                                h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            

    

    document.save('benyehuda_book_from_athuor_'+url.split('/')[-1]+'.docx')

